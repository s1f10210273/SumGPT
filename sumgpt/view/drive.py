#google drive api関連の処理
from django.shortcuts import get_object_or_404
from django.http import HttpResponse, JsonResponse, Http404
from ..models import Sum, Doc
import os
from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive
from urllib.parse import quote
from docx import Document  

def access_drive():
    #Google認証
    current_dir_path = os.path.dirname(os.path.abspath(__file__))
    settings_file_path = os.path.join(current_dir_path,'../pydrive/settings.yaml')
    gauth = GoogleAuth(settings_file=settings_file_path)
    gauth.LocalWebserverAuth()
    return GoogleDrive(gauth)

#Docsの新規作成保存
def create_docs(filename):
    try:
        drive = access_drive()
        docs = drive.CreateFile()
        docs.SetContentFile(get_wordfile_path())
        docs["title"] = filename
        #ファイルのアップロード(google driveの最上位のディレクトリに置かれます。)
        docs.Upload(param={'convert': True})
        #作成したファイルのURLを返す
        return 'https://docs.google.com/document/d/' + str(docs['id'])
    except:
        return "ERROR"

def get_wordfile_path():
    current_dir_path = os.path.dirname(os.path.abspath(__file__))
    settings_file_path = os.path.join(current_dir_path,'../pydrive/output.docx')
    return settings_file_path


#wordファイルの作成関数 戻り値はファイル生成の可否
def create_wordfile(sum:Sum):
    try:
        replace_dic = {"@要約文挿入@": sum.sum, "@原文挿入@" : sum.detail, "@ユーザー名挿入@" : sum.user, "@タイムスタンプ挿入@" : sum.timestamp.strftime('%Y/%m/%d %H:%M:%S')}
        current_dir_path = os.path.dirname(os.path.abspath(__file__))
        static_file_path = os.path.join(current_dir_path,'../static/template.docx')
        doc = Document(static_file_path)
        for i in range(len(doc.paragraphs)):
            if (doc.paragraphs[i].text == ""): continue
            if (doc.paragraphs[i].runs[0].underline == False): continue
            if(doc.paragraphs[i].runs[0].bold): continue
            for k, v in replace_dic.items():
                doc.paragraphs[i].text =  doc.paragraphs[i].text.replace(k, str(v).replace("\r", ""))
        settings_file_path = get_wordfile_path()
        doc.save(settings_file_path)
        return True
    except:
        return False


def docx_download(request):
    if request.method == 'POST':
        filename = request.POST.get("filename", None)
        pk = request.POST.get("pk", None)
        if Sum.objects.filter(pk=pk).exists() == False: 
            return HttpResponse()
        sum = Sum.objects.get(pk=pk)
        if create_wordfile(sum) == False: 
            return HttpResponse()
        doc = Document(get_wordfile_path())
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename*=UTF-8\'\'{}.docx'.format(quote(filename.encode('utf-8')))
        doc.save(response)
        return response
    raise Http404()

def docx_to_drive(request):
    drive_msg = {"msg":"GoogleDriveへの保存に失敗しました。", "url": ""}
    if request.method == 'POST':
        filename = request.POST.get("filename", None)
        pk = request.POST.get("pk", None)

        if Sum.objects.filter(pk=pk).exists() == False: 
            return JsonResponse(drive_msg)
        sum = Sum.objects.get(pk=pk)
        
        if create_wordfile(sum) == False: 
            return JsonResponse(drive_msg)
        docs_url = create_docs(filename)
        if (docs_url != "ERROR"):
            drive_msg["msg"] = "GoogleDriveへの保存に成功しました。"
            drive_msg["url"] = docs_url
            # DBの保存 
            Doc.objects.update_or_create(sum=sum, defaults={"url":docs_url})
        return JsonResponse(drive_msg)
    raise Http404()